from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from docxtpl import DocxTemplate

from app.schemas.templates import TemplateRecord
from app.services.field_registry import flatten_registry, sample_context


STORAGE_ROOT = Path(__file__).resolve().parents[2] / "storage" / "templates"
PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z_][\w]*(?:\.[a-zA-Z_][\w]*)*)\s*\}\}")


class DocxExportError(RuntimeError):
    pass


def template_dir(template_id: str) -> Path:
    path = STORAGE_ROOT / template_id
    path.mkdir(parents=True, exist_ok=True)
    return path


def export_template_docx(template: TemplateRecord) -> Path:
    target = template_dir(template.id) / "template.docx"
    document = Document()
    _configure_document(document)
    content = template.editor_json.get("content") or []
    if not content:
        raise DocxExportError("Template editor_json has no content.")
    for node in content:
        _write_block(document, node)
    document.save(target)
    return target


def render_template_docx(template: TemplateRecord, template_docx_path: Path | None = None) -> Path:
    source = template_docx_path or export_template_docx(template)
    rendered = template_dir(template.id) / "rendered-sample.docx"
    docx_template = DocxTemplate(str(source))
    docx_template.render(sample_context(template.document_type))
    docx_template.save(rendered)
    return rendered


def extract_placeholders_from_docx(docx_path: Path) -> list[str]:
    with zipfile.ZipFile(docx_path) as archive:
        names = [name for name in archive.namelist() if name.startswith("word/") and name.endswith(".xml")]
        text = "\n".join(archive.read(name).decode("utf-8", errors="ignore") for name in names)
    return sorted(set(PLACEHOLDER_RE.findall(text)))


def store_uploaded_docx(template: TemplateRecord, uploaded_path: Path) -> Path:
    target = template_dir(template.id) / "original.docx"
    shutil.copyfile(uploaded_path, target)
    return target


def unknown_placeholders(template: TemplateRecord, placeholders: list[str]) -> list[str]:
    known = set(flatten_registry(template.document_type))
    return sorted(set(placeholders) - known)


def _configure_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(14)


def _write_block(document: DocumentObject, node: dict[str, Any]) -> None:
    node_type = node.get("type")
    if node_type == "paragraph":
        _write_paragraph(document.add_paragraph(), node)
    elif node_type == "heading":
        level = int((node.get("attrs") or {}).get("level") or 1)
        paragraph = document.add_heading(level=min(max(level, 1), 3))
        _write_paragraph(paragraph, node)
    elif node_type == "bulletList":
        for item in node.get("content") or []:
            _write_list_item(document, item, "List Bullet")
    elif node_type == "orderedList":
        for item in node.get("content") or []:
            _write_list_item(document, item, "List Number")
    elif node_type == "table":
        _write_table(document, node)
    elif node_type == "hardBreak":
        paragraph = document.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)


def _write_list_item(document: DocumentObject, node: dict[str, Any], style: str) -> None:
    blocks = node.get("content") or []
    if not blocks:
        document.add_paragraph("", style=style)
        return
    for block in blocks:
        paragraph = document.add_paragraph(style=style)
        _write_paragraph(paragraph, block)


def _write_table(document: DocumentObject, node: dict[str, Any]) -> None:
    rows = [row for row in node.get("content") or [] if row.get("type") == "tableRow"]
    if not rows:
        return
    column_count = max(len(row.get("content") or []) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row_node in enumerate(rows):
        cells = row_node.get("content") or []
        for cell_index in range(column_count):
            cell = table.rows[row_index].cells[cell_index]
            cell_node = cells[cell_index] if cell_index < len(cells) else {"content": []}
            cell.text = ""
            blocks = cell_node.get("content") or [{"type": "paragraph", "content": []}]
            for block_index, block in enumerate(blocks):
                paragraph = cell.paragraphs[0] if block_index == 0 else cell.add_paragraph()
                _write_paragraph(paragraph, block)


def _write_paragraph(paragraph: Any, node: dict[str, Any]) -> None:
    for child in node.get("content") or []:
        child_type = child.get("type")
        if child_type == "text":
            run = paragraph.add_run(str(child.get("text") or ""))
            marks = child.get("marks") or []
            run.bold = any(mark.get("type") == "bold" for mark in marks)
            run.italic = any(mark.get("type") == "italic" for mark in marks)
            run.underline = any(mark.get("type") == "underline" for mark in marks)
        elif child_type == "templateField":
            attrs = child.get("attrs") or {}
            paragraph.add_run(str(attrs.get("jinja") or ""))
        elif child_type == "hardBreak":
            paragraph.add_run().add_break()

